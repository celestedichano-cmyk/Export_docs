#!/usr/bin/env python3
"""
Export Documentation Tool - Flask Backend
"""

import os
import re
import copy
import json
import subprocess
from datetime import datetime
from io import BytesIO
from flask import Flask, request, jsonify, send_file, send_from_directory
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__, static_folder='.')
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates')

MONTHS_ES = {
    1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
    5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
    9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
}

def today_es():
    d = datetime.today()
    return f"{d.day} de {MONTHS_ES[d.month]} del {d.year}"

# Campos de firmante compartidos por todos los templates
FIRMANTE_FIELDS = [
    {"id": "firmante_nombre", "label": "Nombre del firmante", "type": "text", "placeholder": "Mariana Brizzio"},
    {"id": "firmante_cargo", "label": "Cargo del firmante", "type": "text", "placeholder": "Head R&D South Cone"},
]

TEMPLATES = {
    "certificado_codificacion": {
        "label": "Certificado de Codificación de Fecha y Lote",
        "file": "Template_CERTIFICADO_DE_CODIFICACIO_N_DE_FECHA_Y_LOTE.docx",
        "help": "El nombre del producto se puede autocompletar cargando la Ficha Técnica (PDF) o el Dossier (xlsx). El resto de los campos se completa manualmente. Los campos de Estructura y Significado vienen prellenados con el formato estándar — editalos solo si tu caso es distinto.",
        "fields": [
            {"id": "producto", "label": "Nombre del producto", "type": "text", "placeholder": "NOT Burger 113g"},
            {"id": "estructura_elaboracion", "label": "Estructura — Fecha de Elaboración", "type": "text", "placeholder": "DD/MM/AA", "prefill": True},
            {"id": "significado_elaboracion", "label": "Significado — Fecha de Elaboración", "type": "text", "placeholder": "D: día / M: mes / A: año", "prefill": True},
            {"id": "estructura_vencimiento", "label": "Estructura — Fecha de Vencimiento", "type": "text", "placeholder": "DD/MM/AA", "prefill": True},
            {"id": "significado_vencimiento", "label": "Significado — Fecha de Vencimiento", "type": "text", "placeholder": "D: día / M: mes / A: año", "prefill": True},
            {"id": "estructura_lote", "label": "Estructura — Lote", "type": "text", "placeholder": "PB XXX (HH:MM)", "prefill": True},
            {"id": "significado_lote", "label": "Significado — Lote", "type": "textarea", "placeholder": "PB: Identificación copacker Pacificblu\nXXX: Día correlativo del año\n(HH:MM): Hora de envasado", "prefill": True},
            {"id": "fecha", "label": "Fecha del documento", "type": "text", "placeholder": today_es()},
        ] + FIRMANTE_FIELDS,
    },

    "certificado_empaque": {
        "label": "Certificado de Empaque",
        "file": "Template_CERTIFICADO_DE_EMPAQUE.docx",
        "help": "Cargá la Ficha Técnica (PDF) para autocompletar el nombre del producto y las características físicas de envase primario, secundario y caja master. El nombre del producto también se autocompleta con el Dossier (xlsx). Tildá qué secciones incluir en el documento final.",
        "fields": [
            {"id": "producto", "label": "Nombre del producto", "type": "text", "placeholder": "NOT Burger 113g"},
            {"id": "incluir_primario", "label": "Incluir Envase Primario", "type": "checkbox", "checked": True},
            {"id": "incluir_secundario", "label": "Incluir Envase Secundario", "type": "checkbox", "checked": True},
            {"id": "incluir_caja_master", "label": "Incluir Caja Master", "type": "checkbox", "checked": False},
            {"id": "tipo_envase_primario", "label": "Tipo de envase primario", "type": "text", "placeholder": "Bolsa plástica"},
            {"id": "ancho_p", "label": "Primario – Ancho (cm)", "type": "number", "placeholder": ""},
            {"id": "largo_p", "label": "Primario – Largo (cm)", "type": "number", "placeholder": ""},
            {"id": "alto_p", "label": "Primario – Alto (cm)", "type": "number", "placeholder": ""},
            {"id": "peso_neto_p", "label": "Primario – Peso Neto (g)", "type": "number", "placeholder": ""},
            {"id": "peso_bruto_p", "label": "Primario – Peso Bruto (g)", "type": "number", "placeholder": ""},
            {"id": "material_p", "label": "Primario – Material", "type": "text", "placeholder": "PVDC/PE"},
            {"id": "tipo_envase_secundario", "label": "Tipo de envase secundario", "type": "text", "placeholder": "Caja de cartón"},
            {"id": "ancho_s", "label": "Secundario – Ancho (cm)", "type": "number", "placeholder": ""},
            {"id": "largo_s", "label": "Secundario – Largo (cm)", "type": "number", "placeholder": ""},
            {"id": "alto_s", "label": "Secundario – Alto (cm)", "type": "number", "placeholder": ""},
            {"id": "peso_neto_s", "label": "Secundario – Peso Neto (kg)", "type": "number", "placeholder": ""},
            {"id": "peso_bruto_s", "label": "Secundario – Peso Bruto (kg)", "type": "number", "placeholder": ""},
            {"id": "material_s", "label": "Secundario – Material", "type": "text", "placeholder": "Cartón corrugado"},
            {"id": "cantidad_unidades_s", "label": "Secundario – Cantidad de unidades", "type": "number", "placeholder": ""},
            {"id": "tipo_envase_caja_master", "label": "Tipo de envase Caja Master", "type": "text", "placeholder": "Cartón corrugado"},
            {"id": "ancho_cm", "label": "Caja Master – Ancho (cm)", "type": "number", "placeholder": ""},
            {"id": "largo_cm", "label": "Caja Master – Largo (cm)", "type": "number", "placeholder": ""},
            {"id": "alto_cm", "label": "Caja Master – Alto (cm)", "type": "number", "placeholder": ""},
            {"id": "peso_neto_cm", "label": "Caja Master – Peso Neto (kg)", "type": "number", "placeholder": ""},
            {"id": "peso_bruto_cm", "label": "Caja Master – Peso Bruto (kg)", "type": "number", "placeholder": ""},
            {"id": "material_cm", "label": "Caja Master – Material", "type": "text", "placeholder": "Cartón corrugado"},
            {"id": "cantidad_unidades", "label": "Caja Master – Cantidad de unidades", "type": "number", "placeholder": ""},
        ] + FIRMANTE_FIELDS,
    },

    "certificado_proceso": {
        "label": "Certificado de Proceso de Producción",
        "file": "Template_CERTIFICADO_PROCESO_DE_PRODUCCIO_N.docx",
        "help": "El nombre del producto se puede autocompletar cargando la Ficha Técnica (PDF) o el Dossier (xlsx). Además, subí una imagen (.png/.jpg) del diagrama de flujo, exportada manualmente desde el PPTX de procesos de producción.",
        "fields": [
            {"id": "producto", "label": "Nombre del producto (código)", "type": "text", "placeholder": "NOT20012"},
            {"id": "fecha", "label": "Fecha del documento", "type": "text", "placeholder": today_es()},
        ] + FIRMANTE_FIELDS,
    },

    "informe_analisis": {
        "label": "Informe de Análisis",
        "file": "Template_INFORME_ANA_LISIS.docx",
        "help": "Cargá la Ficha Técnica (PDF) para autocompletar el nombre del producto y los parámetros sensoriales, fisicoquímicos y microbiológicos. El nombre del producto también se autocompleta con el Dossier (xlsx). También podés cargar un Informe de Análisis modelo o anterior (.docx) para autocompletar automáticamente la columna de Metodología en Fisicoquímicos y Microbiológicos, reutilizando las metodologías ya usadas en ese informe sin tener que volver a tipearlas. Tildá qué tablas incluir en el documento final.",
        "fields": [
            {"id": "producto", "label": "Nombre del producto", "type": "text", "placeholder": "NOT Burger 113g"},
            {"id": "incluir_fq", "label": "Incluir Análisis Fisicoquímicos", "type": "checkbox", "checked": True},
            {"id": "incluir_mb", "label": "Incluir Análisis Microbiológicos", "type": "checkbox", "checked": True},
            {"id": "incluir_sensorial", "label": "Incluir Análisis Sensoriales", "type": "checkbox", "checked": True},
            {"id": "incluir_contaminantes", "label": "Incluir Límites de Contaminantes", "type": "checkbox", "checked": True},
            {"id": "fq_rows", "label": "Análisis fisicoquímicos (Parámetro | Metodología | Resultado, una fila por línea)", "type": "textarea", "placeholder": "pH | AOAC 943.02 | 6,8\nHumedad | AOAC 925.10 | 62%"},
            {"id": "mb_rows", "label": "Análisis microbiológicos (Parámetro | Metodología | Resultado)", "type": "textarea", "placeholder": "Recuento aeróbico | ISO 4833 | <10 UFC/g"},
            {"id": "apariencia", "label": "Sensorial – Apariencia (dejar vacío si no aplica)", "type": "text", "placeholder": ""},
            {"id": "color", "label": "Sensorial – Color", "type": "text", "placeholder": "Marrón característico"},
            {"id": "aroma", "label": "Sensorial – Aroma", "type": "text", "placeholder": "Característico a carne"},
            {"id": "sabor", "label": "Sensorial – Sabor", "type": "text", "placeholder": "Umami, levemente salado"},
            {"id": "textura", "label": "Sensorial – Textura", "type": "text", "placeholder": "Firme, jugosa"},
            {"id": "pb", "label": "Contaminantes – Plomo Pb (mg/kg)", "type": "text", "placeholder": "<0,1"},
            {"id": "cu", "label": "Contaminantes – Cobre Cu (mg/kg)", "type": "text", "placeholder": "<0,5"},
            {"id": "as_", "label": "Contaminantes – Arsénico As (mg/kg)", "type": "text", "placeholder": "<0,1"},
            {"id": "sn", "label": "Contaminantes – Estaño Sn (mg/kg)", "type": "text", "placeholder": "<1,0"},
            {"id": "fe", "label": "Contaminantes – Hierro Fe (mg/kg)", "type": "text", "placeholder": "<5,0"},
            {"id": "fecha", "label": "Fecha del documento", "type": "text", "placeholder": today_es()},
        ] + FIRMANTE_FIELDS,
    },

    "informe_nutricional": {
        "label": "Informe de Análisis Nutricional",
        "file": "Template_INFORME_ANA_LISIS_NUTRICIONAL.docx",
        "help": "Cargá la Ficha Técnica (PDF) o el Dossier (xlsx) para autocompletar el nombre del producto y la tabla de información nutricional (columna 100g/100ml). En el Dossier, estos datos se toman de la hoja 'Proyecto de rótulo', sección Información Nutricional, columna 100 ml.",
        "fields": [
            {"id": "producto", "label": "Nombre del producto", "type": "text", "placeholder": "NOT Burger 113g"},
            {"id": "energia", "label": "Energía (kcal)", "type": "number", "placeholder": ""},
            {"id": "proteina", "label": "Proteína (g)", "type": "number", "placeholder": ""},
            {"id": "grasa_total", "label": "Grasa total (g)", "type": "number", "placeholder": ""},
            {"id": "grasa_sat", "label": "Grasa saturada (g)", "type": "number", "placeholder": ""},
            {"id": "grasa_mono", "label": "Grasa monoinsaturada (g)", "type": "number", "placeholder": ""},
            {"id": "grasa_poli", "label": "Grasa poliinsaturada (g)", "type": "number", "placeholder": ""},
            {"id": "grasa_trans", "label": "Grasa trans (g)", "type": "number", "placeholder": ""},
            {"id": "colesterol", "label": "Colesterol (mg)", "type": "number", "placeholder": ""},
            {"id": "carb_totales", "label": "Carbohidratos totales (g)", "type": "number", "placeholder": ""},
            {"id": "carb_disp", "label": "Carbohidratos disponibles (g)", "type": "number", "placeholder": ""},
            {"id": "azucares", "label": "Azúcares (g)", "type": "number", "placeholder": ""},
            {"id": "fibra", "label": "Fibra (g)", "type": "number", "placeholder": ""},
            {"id": "sodio", "label": "Sodio (mg)", "type": "number", "placeholder": ""},
            {"id": "micronutrientes_extra", "label": "Micronutrientes adicionales (Nombre (unidad) | Valor, una fila por línea — ej: Calcio (mg) | 129)", "type": "textarea", "placeholder": ""},
            {"id": "fecha", "label": "Fecha del documento", "type": "text", "placeholder": today_es()},
        ] + FIRMANTE_FIELDS,
    },

    "informe_aditivos": {
        "label": "Informe de Funcionalidad de Aditivos",
        "file": "Template_INFORME_FUNCIONALIDAD_ADITIVOS.docx",
        "help": "Cargá la Ficha Técnica (PDF) para autocompletar el nombre del producto y traer el listado de ingredientes de la sección Descripción del producto, o cargá el Dossier (xlsx) para autocompletar el nombre del producto y traer la tabla de ingredientes en orden decreciente (hoja 'Fórmula'). Solo el Dossier indica además cuáles son aditivos y completa su función tecnológica y su INS, ambos copiados directamente del Dossier; si alguno de los dos no figura ahí, queda sin completar. Con la FT esos datos hay que completarlos manualmente.",
        "fields": [
            {"id": "producto", "label": "Nombre del producto", "type": "text", "placeholder": "NOT Burger 113g"},
            {"id": "ingredientes", "label": "Lista de ingredientes (uno por línea)", "type": "textarea", "placeholder": "Agua\nProteína de soya\nAceite de girasol\nSal\nAromati­zante natural"},
            {"id": "aditivos_filas", "label": "Nº de filas que son aditivos (ej: 3, 5, 8)", "type": "text", "placeholder": "3, 5"},
            {"id": "aditivos", "label": "Aditivos y su función (Nombre | Función, uno por línea)", "type": "textarea", "placeholder": "ACEITE DE GIRASOL | Agente de relleno\nAROMATIZANTE NATURAL | Aromatizante"},
            {"id": "fecha", "label": "Fecha del documento", "type": "text", "placeholder": today_es()},
        ] + FIRMANTE_FIELDS,
    },

    "reporte_fibra": {
        "label": "Reporte de Fibra Dietética Total",
        "file": "Template_REPORTE_FIBRA_DIETE_TICA_TOTAL.docx",
        "help": "Cargá la Ficha Técnica (PDF) para autocompletar el nombre del producto y la Fibra Dietética Total (columna 100g). El nombre del producto también se autocompleta con el Dossier (xlsx).",
        "fields": [
            {"id": "producto", "label": "Nombre del producto", "type": "text", "placeholder": "NOT Burger 113g"},
            {"id": "fibra_total", "label": "Fibra Dietética Total (g/100g)", "type": "number", "placeholder": ""},
            {"id": "fecha", "label": "Fecha del documento", "type": "text", "placeholder": today_es()},
        ] + FIRMANTE_FIELDS,
    },

    "reporte_formula": {
        "label": "Reporte de Fórmula de Producto",
        "file": "Template_REPORTE_FO_RMULA.docx",
        "help": "Cargá el Dossier (xlsx) para autocompletar el nombre del producto y la lista de ingredientes con su % en orden decreciente (hoja 'Fórmula', tabla de orden decreciente). El nombre del producto también se autocompleta con la Ficha Técnica (PDF).",
        "fields": [
            {"id": "producto", "label": "Nombre del producto", "type": "text", "placeholder": "NOT Burger 113g"},
            {"id": "formula_rows", "label": "Ingredientes y cantidades (Ingrediente | % en 100g, uno por línea)", "type": "textarea", "placeholder": "Agua | 55,0\nProteína de soya texturizada | 20,0\nAceite de girasol | 10,0"},
            {"id": "fecha", "label": "Fecha del documento", "type": "text", "placeholder": today_es()},
        ] + FIRMANTE_FIELDS,
    },

    "reporte_saborizantes": {
        "label": "Reporte de Saborizantes",
        "file": "Template__REPORTE_SABORIZANTES.docx",
        "help": "El nombre del producto se puede autocompletar cargando la Ficha Técnica (PDF) o el Dossier (xlsx). El resto de los campos se completa manualmente. El campo Total se calcula solo a partir del % Natural y % Idéntico Natural. Si dejás vacío Natural o Idéntico Natural, esa fila se elimina del documento generado en vez de usar un valor de ejemplo.",
        "fields": [
            {"id": "producto", "label": "Nombre del producto", "type": "text", "placeholder": "NOT Burger 113g"},
            {"id": "sab_natural", "label": "Saborizante Natural (%)", "type": "text", "placeholder": "0,5", "sin_relleno_de_ejemplo": True},
            {"id": "sab_identico", "label": "Saborizante Idéntico Natural (%)", "type": "text", "placeholder": "0,3", "sin_relleno_de_ejemplo": True},
            {"id": "sab_total", "label": "Total saborizantes (%)", "type": "text", "placeholder": "0,8", "sin_relleno_de_ejemplo": True},
            {"id": "fecha", "label": "Fecha del documento", "type": "text", "placeholder": today_es()},
        ] + FIRMANTE_FIELDS,
    },
}


# ─── Document generation helpers ─────────────────────────────────────────────

def replace_in_paragraph(para, replacements):
    """
    Replace text preserving per-run formatting (font, bold, size etc.).
    Tries per-run replacement first; falls back to collapsing runs only
    when a placeholder spans multiple runs.
    """
    full = ''.join(r.text for r in para.runs)
    if not any(old in full for old in replacements):
        return
    # Per-run replacement (preserves formatting)
    for run in para.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)
    # Fallback: if placeholder still present it spans multiple runs
    full_after = ''.join(r.text for r in para.runs)
    for old, new in replacements.items():
        if old in full_after:
            collapsed = full_after.replace(old, new)
            if para.runs:
                para.runs[0].text = collapsed
                for r in para.runs[1:]:
                    r.text = ''
            break

def replace_all(doc, replacements):
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            if hdr:
                for para in hdr.paragraphs:
                    replace_in_paragraph(para, replacements)

def set_producto_bold(doc, producto):
    """After replacing XXX with producto, find those runs and force bold=True."""
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text == producto:
                run.bold = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text == producto:
                            run.bold = True

def apply_firmante(doc, data):
    """Replace firmante name and cargo in all paragraphs."""
    nombre = data.get('firmante_nombre', '').strip()
    cargo = data.get('firmante_cargo', '').strip()
    if nombre:
        replace_all(doc, {'Mariana Brizzio': nombre, 'Javiera Mujica': nombre})
    if cargo:
        replace_all(doc, {
            'Head R&D South Cone': cargo,
            'Head R****&****D South Cone': cargo,
            'R&D Manager South Cone': cargo,
            'R****&****D Manager South Cone': cargo,
        })

def fill_table_rows(table, data_rows, start_row=1):
    while len(table.rows) > start_row:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    WNS_local = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for row_data in data_rows:
        row = copy.deepcopy(table.rows[start_row - 1])
        cells = row.cells
        for i, val in enumerate(row_data):
            if i < len(cells):
                for para in cells[i].paragraphs:
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = str(val)
                        para.runs[0].font.name = 'NotFont Display'
                    else:
                        run = para.add_run(str(val))
                        run.font.name = 'NotFont Display'
        # Forzar shading blanco para anular el banding heredado del estilo
        # de tabla (Table1/TableNormal), que de otro modo pinta filas de gris
        # de forma alternada al clonar filas dinámicamente.
        for cell in cells:
            tcPr = cell._tc.find(f'{{{WNS_local}}}tcPr')
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                cell._tc.insert(0, tcPr)
            shd = tcPr.find(f'{{{WNS_local}}}shd')
            if shd is None:
                shd = OxmlElement('w:shd')
                tcPr.append(shd)
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'ffffff')
        table._tbl.append(row._tr)

def parse_rows(text, separator='|'):
    rows = []
    for line in text.strip().split('\n'):
        if line.strip():
            parts = [p.strip() for p in line.split(separator)]
            rows.append(parts)
    return rows

def set_cell_value(cell, val, font_name='NotFont Display', font_size=127000):
    """Write a value into a cell preserving NotFont Display typography."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = str(val)
            p.runs[0].font.name = font_name
            p.runs[0].font.size = font_size
        else:
            run = p.add_run(str(val))
            run.font.name = font_name
            run.font.size = font_size


# ─── Per-template generation ─────────────────────────────────────────────

def set_cell_value_xml(cell_el, val, WNS):
    """Write value into an lxml cell element with NotFont Display font."""
    runs = cell_el.findall(f'.//{{{WNS}}}r')
    if runs:
        for r in runs:
            for t_el in r.findall(f'{{{WNS}}}t'):
                t_el.text = ''
        t_el = runs[0].find(f'{{{WNS}}}t')
        if t_el is None:
            t_el = OxmlElement('w:t')
            runs[0].append(t_el)
        t_el.text = val
        rPr = runs[0].find(f'{{{WNS}}}rPr')
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            runs[0].insert(0, rPr)
        rFonts = rPr.find(f'{{{WNS}}}rFonts')
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(f'{{{WNS}}}ascii', 'NotFont Display')
        rFonts.set(f'{{{WNS}}}hAnsi', 'NotFont Display')
    else:
        paras = cell_el.findall(f'.//{{{WNS}}}p')
        if paras:
            r_new = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(f'{{{WNS}}}ascii', 'NotFont Display')
            rFonts.set(f'{{{WNS}}}hAnsi', 'NotFont Display')
            rPr.append(rFonts)
            r_new.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.text = val
            r_new.append(t_el)
            paras[0].append(r_new)

def generate_doc(template_id, data):
    tmpl = TEMPLATES[template_id]
    path = os.path.join(TEMPLATES_DIR, tmpl['file'])
    doc = Document(path)

    producto = data.get('producto', 'XXX').upper()
    fecha = data.get('fecha', today_es())

    if template_id == 'certificado_codificacion':
        if doc.tables:
            t = doc.tables[0]

            def set_row_text(row, col_idx, text):
                cell = row.cells[col_idx]
                lines = text.split('\n')
                paras = cell.paragraphs
                # Si hay más líneas que párrafos disponibles, usamos el último
                # párrafo para las líneas sobrantes (separadas por salto de línea
                # dentro del mismo run, ya que python-docx no crea párrafos nuevos
                # automáticamente en una celda existente sin más estructura).
                if len(lines) <= len(paras):
                    for i, p in enumerate(paras):
                        for r in p.runs:
                            r.text = ''
                        val = lines[i] if i < len(lines) else ''
                        if p.runs:
                            p.runs[0].text = val
                            p.runs[0].font.name = 'NotFont Display'
                        elif val:
                            run = p.add_run(val)
                            run.font.name = 'NotFont Display'
                else:
                    # Más líneas que párrafos: juntamos todo en el primer párrafo
                    first_p = paras[0]
                    for p in paras:
                        for r in p.runs:
                            r.text = ''
                    if first_p.runs:
                        first_p.runs[0].text = lines[0]
                        first_p.runs[0].font.name = 'NotFont Display'
                    for extra_line in lines[1:]:
                        run = first_p.add_run('\n' + extra_line)
                        run.font.name = 'NotFont Display'

            # Fila 1: Fecha de Elaboración (Estructura=col1, Significado=col2)
            est_elab = data.get('estructura_elaboracion', '').strip()
            if est_elab and est_elab != 'DD/MM/AA':
                set_row_text(t.rows[1], 1, est_elab)
            sig_elab = data.get('significado_elaboracion', '').strip()
            if sig_elab and sig_elab != 'D: día / M: mes / A: año':
                set_row_text(t.rows[1], 2, sig_elab)

            # Fila 2: Fecha de Vencimiento
            est_venc = data.get('estructura_vencimiento', '').strip()
            if est_venc and est_venc != 'DD/MM/AA':
                set_row_text(t.rows[2], 1, est_venc)
            sig_venc = data.get('significado_vencimiento', '').strip()
            if sig_venc and sig_venc != 'D: día / M: mes / A: año':
                set_row_text(t.rows[2], 2, sig_venc)

            # Fila 3: Lote — mismo comportamiento: si no se edita, queda el
            # texto crudo del template tal cual (PB XXX (HH:MM) / Pacificblu)
            default_sig_lote = 'PB: Identificación copacker Pacificblu\nXXX: Día correlativo del año\n(HH:MM): Hora de envasado'
            est_lote = data.get('estructura_lote', '').strip()
            if est_lote and est_lote != 'PB XXX (HH:MM)':
                set_row_text(t.rows[3], 1, est_lote)
            sig_lote = data.get('significado_lote', '').strip()
            if sig_lote and sig_lote != default_sig_lote:
                set_row_text(t.rows[3], 2, sig_lote)

        for para in doc.paragraphs:
            replace_in_paragraph(para, {'XXX': producto, 'XX de XX del 2025': fecha})

    elif template_id == 'certificado_proceso':
        replace_all(doc, {'NOTXXX': producto, 'XX de mayo de 2025': fecha})
        img_b64 = data.get('imagen_proceso_b64')
        if img_b64:
            import base64
            img_bytes = base64.b64decode(img_b64)
            # Find the blank paragraph between the intro text and "Se extiende..."
            paras = doc.paragraphs
            intro_idx = None
            extiende_idx = None
            for i, p in enumerate(paras):
                if 'diagrama de flujo' in p.text.lower():
                    intro_idx = i
                if 'se extiende el presente' in p.text.lower():
                    extiende_idx = i
                    break
            if intro_idx is not None and extiende_idx is not None and extiende_idx > intro_idx + 1:
                # Use the first blank paragraph right after the intro text
                target_p = paras[intro_idx + 1]
                run = target_p.add_run()
                run.add_picture(BytesIO(img_bytes), width=Inches(6))

    elif template_id == 'certificado_empaque':
        incluir_primario = data.get('incluir_primario', 'true') == 'true'
        incluir_secundario = data.get('incluir_secundario', 'true') == 'true'
        incluir_caja = data.get('incluir_caja_master', 'false') == 'true'

        # Replace tipo de envase BEFORE replace_all so 'XXX' isn't clobbered by producto.
        # Orden de aparición en el template: Primario, Secundario, Caja Master.
        found = 0
        for para in doc.paragraphs:
            if 'Tipo de envase:' in para.text:
                found += 1
                if found == 1:
                    replace_in_paragraph(para, {'Tipo de envase: XXX': f'Tipo de envase: {data.get("tipo_envase_primario","")}'})
                elif found == 2:
                    replace_in_paragraph(para, {'Tipo de envase: XXX': f'Tipo de envase: {data.get("tipo_envase_secundario","")}'})
                elif found == 3:
                    replace_in_paragraph(para, {'Tipo de envase: XXX': f'Tipo de envase: {data.get("tipo_envase_caja_master","Cartón corrugado")}'})
        replace_all(doc, {'XXX': producto, 'XX de XX del 2025': fecha})

        tables = doc.tables
        if len(tables) >= 1:
            mapping_p = [
                ('Ancho (cm)', data.get('ancho_p','')),
                ('Largo (cm)', data.get('largo_p','')),
                ('Alto (cm)', data.get('alto_p','')),
                ('Peso Neto (g)', data.get('peso_neto_p','')),
                ('Peso Bruto (g)', data.get('peso_bruto_p','')),
                ('Material', data.get('material_p','')),
            ]
            for row in tables[0].rows[1:]:
                label = row.cells[0].text.strip()
                for lbl, val in mapping_p:
                    if lbl in label:
                        set_cell_value(row.cells[1], val)
        if len(tables) >= 2:
            mapping_s = [
                ('Ancho (cm)', data.get('ancho_s','')),
                ('Largo (cm)', data.get('largo_s','')),
                ('Alto (cm)', data.get('alto_s','')),
                ('Peso Neto (kg)', data.get('peso_neto_s','')),
                ('Peso Bruto (kg)', data.get('peso_bruto_s','')),
                ('Material', data.get('material_s','')),
                ('Cantidad de unidades', data.get('cantidad_unidades_s','')),
            ]
            for row in tables[1].rows[1:]:
                label = row.cells[0].text.strip()
                for lbl, val in mapping_s:
                    if lbl in label:
                        set_cell_value(row.cells[1], val)
        if len(tables) >= 3:
            mapping_cm = [
                ('Ancho (cm)', data.get('ancho_cm','')),
                ('Largo (cm)', data.get('largo_cm','')),
                ('Alto (cm)', data.get('alto_cm','')),
                ('Peso Neto (kg)', data.get('peso_neto_cm','')),
                ('Peso Bruto (kg)', data.get('peso_bruto_cm','')),
                ('Material', data.get('material_cm','')),
                ('Cantidad de unidades', data.get('cantidad_unidades','')),
            ]
            for row in tables[2].rows[1:]:
                label = row.cells[0].text.strip()
                for lbl, val in mapping_cm:
                    if lbl in label:
                        set_cell_value(row.cells[1], val)

        # Eliminar secciones no marcadas. Cada sección = 1 párrafo título +
        # 1 párrafo "Tipo de envase" + 1 tabla, en orden Primario/Secundario/
        # Caja Master. Se eliminan en orden inverso para no desordenar índices
        # mientras se itera sobre el body.
        secciones_incluir = [incluir_primario, incluir_secundario, incluir_caja]
        WNS_local = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        body = doc.element.body

        def encontrar_secciones():
            """Detecta cada bloque (título, tipo_envase_p, tabla) en orden de aparición."""
            bloques = []
            children = list(body)
            i = 0
            while i < len(children):
                el = children[i]
                if el.tag == f'{{{WNS_local}}}p':
                    texto = ''.join(t.text or '' for t in el.findall(f'.//{{{WNS_local}}}t'))
                    if texto.strip() in ('Envase Primario', 'Envase secundario', 'Caja Master'):
                        titulo_el = el
                        tipo_envase_el = children[i + 1] if i + 1 < len(children) else None
                        tabla_el = children[i + 2] if i + 2 < len(children) else None
                        bloques.append((titulo_el, tipo_envase_el, tabla_el))
                        i += 3
                        continue
                i += 1
            return bloques

        bloques = encontrar_secciones()
        for incluir, bloque in zip(secciones_incluir, bloques):
            if not incluir:
                for el in bloque:
                    if el is not None and el.getparent() is not None:
                        el.getparent().remove(el)

    elif template_id == 'informe_analisis':
        replace_all(doc, {'XXX': producto, 'XX de XX del 2025': fecha})
        tables = doc.tables
        if len(tables) > 0 and data.get('fq_rows'):
            fill_table_rows(tables[0], parse_rows(data['fq_rows']))
        if len(tables) > 1 and data.get('mb_rows'):
            fill_table_rows(tables[1], parse_rows(data['mb_rows']))
        # Tables 2 (sensorial) and 3 (contaminantes) are in sdtContent — use xpath
        WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        all_tbls = doc.element.body.findall(f'.//{{{WNS}}}tbl')
        sensorial = {
            'Apariencia': data.get('apariencia',''),
            'Color': data.get('color',''),
            'Aroma': data.get('aroma',''),
            'Sabor': data.get('sabor',''),
            'Textura': data.get('textura',''),
        }
        if len(all_tbls) > 2:
            for row in list(all_tbls[2].findall(f'{{{WNS}}}tr'))[1:]:
                cells = row.findall(f'{{{WNS}}}tc')
                if not cells: continue
                key = ''.join(t.text for t in cells[0].findall(f'.//{{{WNS}}}t') if t.text).strip()
                if key in sensorial and len(cells) > 2:
                    if not sensorial[key].strip():
                        # Sin valor (ej. Apariencia no presente en la FT) —
                        # se elimina la fila en lugar de dejarla vacía.
                        row.getparent().remove(row)
                        continue
                    set_cell_value_xml(cells[2], sensorial[key], WNS)
        contam = {
            'Plomo (Pb)': data.get('pb',''),
            'Cobre (Cu)': data.get('cu',''),
            'Arsénico (As)': data.get('as_',''),
            'Estaño (Sn)': data.get('sn',''),
            'Hierro (Fe)': data.get('fe',''),
        }
        if len(all_tbls) > 3:
            for row in all_tbls[3].findall(f'{{{WNS}}}tr')[1:]:
                cells = row.findall(f'{{{WNS}}}tc')
                if not cells: continue
                key = ''.join(t.text for t in cells[0].findall(f'.//{{{WNS}}}t') if t.text).strip()
                for k, v in contam.items():
                    if k in key and len(cells) > 2:
                        set_cell_value_xml(cells[2], str(v), WNS)

        # Eliminar secciones no marcadas. Cada sección = 1 párrafo título +
        # 1 bloque de tabla (tbl directo, o sdt que contiene una tbl), en
        # orden: Fisicoquímicos, Microbiológicos, Sensoriales, Contaminantes.
        incluir_fq = data.get('incluir_fq', 'true') == 'true'
        incluir_mb = data.get('incluir_mb', 'true') == 'true'
        incluir_sens = data.get('incluir_sensorial', 'true') == 'true'
        incluir_contam = data.get('incluir_contaminantes', 'true') == 'true'
        secciones_incluir = [incluir_fq, incluir_mb, incluir_sens, incluir_contam]

        titulos_seccion = (
            'Análisis fisicoquímicos realizados',
            'Análisis microbiológicos realizados',
            'Análisis sensoriales realizados',
            'Límites de contaminantes',
        )

        def encontrar_bloques_analisis():
            """Detecta cada bloque (párrafo título, elemento de tabla) en orden.
            El elemento de tabla puede ser un <w:tbl> directo o un <w:sdt> que
            contiene una tabla adentro (sdtContent)."""
            bloques = []
            children = list(body)
            i = 0
            while i < len(children):
                el = children[i]
                if el.tag == f'{{{WNS}}}p':
                    texto = ''.join(t.text or '' for t in el.findall(f'.//{{{WNS}}}t'))
                    if any(texto.strip().startswith(tit) for tit in titulos_seccion):
                        titulo_el = el
                        tabla_el = children[i + 1] if i + 1 < len(children) else None
                        bloques.append((titulo_el, tabla_el))
                        i += 2
                        continue
                i += 1
            return bloques

        body = doc.element.body
        bloques_analisis = encontrar_bloques_analisis()
        for incluir, bloque in zip(secciones_incluir, bloques_analisis):
            if not incluir:
                for el in bloque:
                    if el is not None and el.getparent() is not None:
                        el.getparent().remove(el)

    elif template_id == 'informe_nutricional':
        replace_all(doc, {'XXX': producto, 'XX de XX del 2025': fecha})
        nut_map = {
            'Energía (kcal)': 'energia',
            'Proteína (g)': 'proteina',
            'Grasa total (g)': 'grasa_total',
            'Grasa saturada (g)': 'grasa_sat',
            'Grasa monoinsaturada (g)': 'grasa_mono',
            'Grasa poliinsaturada (g)': 'grasa_poli',
            'Grasa trans (g)': 'grasa_trans',
            'Colesterol (mg)': 'colesterol',
            'Carbohidratos totales (g)': 'carb_totales',
            'Carbohidratos disponibles (g)': 'carb_disp',
            'Azúcares (g)': 'azucares',
            'Fibra (g)': 'fibra',
            'Sodio (mg)': 'sodio',
        }
        # Table is inside sdtContent — access via lxml xpath
        WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        nut_tbls = doc.element.body.findall(f'.//{{{WNS}}}tbl')
        if nut_tbls:
            nut_tbl = nut_tbls[0]
            for row in nut_tbl.findall(f'{{{WNS}}}tr')[1:]:
                cells = row.findall(f'{{{WNS}}}tc')
                if len(cells) < 2:
                    continue
                # Get label from col[0]
                key = ''.join(t.text for t in cells[0].findall(f'.//{{{WNS}}}t') if t.text).strip()
                for lbl, fid in nut_map.items():
                    if lbl.strip() in key:
                        val = str(data.get(fid, ''))
                        # Write into col[1] preserving/setting font
                        cell1 = cells[1]
                        # Find existing run or create one
                        runs = cell1.findall(f'.//{{{WNS}}}r')
                        if runs:
                            # Clear all runs then set first
                            for r in runs:
                                for t_el in r.findall(f'{{{WNS}}}t'):
                                    t_el.text = ''
                            t_el = runs[0].find(f'{{{WNS}}}t')
                            if t_el is None:
                                t_el = OxmlElement('w:t')
                                runs[0].append(t_el)
                            t_el.text = val
                            # Force font
                            rPr = runs[0].find(f'{{{WNS}}}rPr')
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                                runs[0].insert(0, rPr)
                            rFonts = rPr.find(f'{{{WNS}}}rFonts')
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rPr.insert(0, rFonts)
                            rFonts.set(f'{{{WNS}}}ascii', 'NotFont Display')
                            rFonts.set(f'{{{WNS}}}hAnsi', 'NotFont Display')
                        else:
                            set_cell_value_xml(cell1, val, WNS)

            # Micronutrientes adicionales: se insertan como filas nuevas al
            # final de la tabla, clonando el formato de la última fila fija
            # (Sodio) para mantener tipografía y bordes consistentes. Como
            # esa fila tiene el borde inferior negro (cierre visual de la
            # tabla), hay que dejarlo solo en la fila realmente última y
            # poner blanco en las intermedias para no duplicar líneas negras.
            micro_raw = data.get('micronutrientes_extra', '').strip()
            if micro_raw:
                filas_actuales = nut_tbl.findall(f'{{{WNS}}}tr')
                if filas_actuales:
                    fila_modelo = filas_actuales[-1]
                    lineas_validas = [l.strip() for l in micro_raw.split('\n') if l.strip() and '|' in l]
                    nuevas_filas = []
                    for linea in lineas_validas:
                        nombre, valor = [p.strip() for p in linea.split('|', 1)]
                        nueva_fila = copy.deepcopy(fila_modelo)
                        celdas_nuevas = nueva_fila.findall(f'{{{WNS}}}tc')
                        if len(celdas_nuevas) >= 2:
                            set_cell_value_xml(celdas_nuevas[0], nombre, WNS)
                            set_cell_value_xml(celdas_nuevas[1], valor, WNS)
                        nut_tbl.append(nueva_fila)
                        nuevas_filas.append(nueva_fila)
                    # Poner borde inferior blanco en Sodio (la fila modelo
                    # original) y en todas las filas nuevas menos la última,
                    # que es la única que debe conservar el cierre negro.
                    filas_a_blanquear = [fila_modelo] + nuevas_filas[:-1]
                    for fila in filas_a_blanquear:
                        for celda in fila.findall(f'{{{WNS}}}tc'):
                            tcPr = celda.find(f'{{{WNS}}}tcPr')
                            if tcPr is None:
                                continue
                            borders = tcPr.find(f'{{{WNS}}}tcBorders')
                            if borders is None:
                                continue
                            bottom = borders.find(f'{{{WNS}}}bottom')
                            if bottom is not None:
                                bottom.set(qn('w:color'), 'ffffff')

    elif template_id == 'informe_aditivos':
        replace_all(doc, {'XX de XX del 2025': fecha})
        replace_all(doc, {'XX': producto})
        if data.get('ingredientes'):
            lines = [l.strip() for l in data['ingredientes'].split('\n') if l.strip()]
            # Parse which row numbers are aditivos (1-based)
            aditivos_filas = set()
            for part in data.get('aditivos_filas', '').split(','):
                part = part.strip()
                if part.isdigit():
                    aditivos_filas.add(int(part))
            # Table is inside sdtContent — find via lxml xpath
            WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            tbls = doc.element.body.findall(f'.//{{{WNS}}}tbl')
            if tbls:
                tbl_el = tbls[0]
                # Remove all existing rows
                for tr in tbl_el.findall(f'{{{WNS}}}tr'):
                    tbl_el.remove(tr)
                for idx, ing in enumerate(lines, start=1):
                    is_aditivo = idx in aditivos_filas
                    texto = ing.upper()
                    tr = OxmlElement('w:tr')
                    trPr = OxmlElement('w:trPr')
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), '315')
                    trHeight.set(qn('w:hRule'), 'atLeast')
                    trPr.append(trHeight)
                    tr.append(trPr)
                    tc = OxmlElement('w:tc')
                    tcPr = OxmlElement('w:tcPr')
                    # Borders: top=gray, others=white (matches template)
                    tcBorders = OxmlElement('w:tcBorders')
                    for side, color in [('top','cccccc'),('left','ffffff'),('bottom','ffffff'),('right','ffffff')]:
                        b = OxmlElement(f'w:{side}')
                        b.set(qn('w:val'), 'single')
                        b.set(qn('w:sz'), '4')
                        b.set(qn('w:space'), '0')
                        b.set(qn('w:color'), color)
                        tcBorders.append(b)
                    tcPr.append(tcBorders)
                    # Cell margins
                    tcMar = OxmlElement('w:tcMar')
                    for side in ['top','left','bottom','right']:
                        m = OxmlElement(f'w:{side}')
                        m.set(qn('w:w'), '40')
                        m.set(qn('w:type'), 'dxa')
                        tcMar.append(m)
                    tcPr.append(tcMar)
                    if is_aditivo:
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'f3f3f3')
                        tcPr.append(shd)
                    tc.append(tcPr)
                    p = OxmlElement('w:p')
                    pPr_ing = OxmlElement('w:pPr')
                    jc_ing = OxmlElement('w:jc')
                    jc_ing.set(qn('w:val'), 'center')
                    pPr_ing.append(jc_ing)
                    p.append(pPr_ing)
                    r = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'NotFont Display')
                    rFonts.set(qn('w:hAnsi'), 'NotFont Display')
                    rPr.append(rFonts)
                    if is_aditivo:
                        rPr.append(OxmlElement('w:b'))
                        rPr.append(OxmlElement('w:bCs'))
                    r.append(rPr)
                    t_el = OxmlElement('w:t')
                    t_el.text = texto
                    r.append(t_el)
                    p.append(r)
                    tc.append(p)
                    tr.append(tc)
                    tbl_el.append(tr)
        if data.get('aditivos'):
            pairs = parse_rows(data['aditivos'])
            aditivo_paras = [p for p in doc.paragraphs if 'ADITIVO' in p.text]

            def fill_aditivo_para(para, nombre, funcion):
                if len(para.runs) >= 2:
                    para.runs[0].text = nombre
                    para.runs[0].bold = True
                    para.runs[1].text = '\n' + funcion
                    para.runs[1].bold = False
                else:
                    replace_in_paragraph(para, {'ADITIVO': nombre, 'Función tecnológica': funcion})
                # Fix justification: "both" causes character spacing on short text
                pPr = para._p.find(qn('w:pPr'))
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None:
                        jc.set(qn('w:val'), 'left')

            # Fill existing fixed placeholders first
            for i, para in enumerate(aditivo_paras):
                if i < len(pairs):
                    nombre = pairs[i][0].upper() if len(pairs[i]) > 0 else ''
                    funcion = pairs[i][1] if len(pairs[i]) > 1 else ''
                    fill_aditivo_para(para, nombre, funcion)

            # If there are more aditivos than placeholders, clone the last
            # placeholder paragraph (with its spacer) for each extra item
            if len(pairs) > len(aditivo_paras) and aditivo_paras:
                last_para = aditivo_paras[-1]
                # Spacer paragraph right after the last ADITIVO block (if blank)
                insert_after_el = last_para._p
                next_p = last_para._p.getnext()
                if next_p is not None and next_p.tag == qn('w:p'):
                    # Treat the immediate following paragraph as a spacer to clone too
                    spacer_el = next_p
                else:
                    spacer_el = None

                for j in range(len(aditivo_paras), len(pairs)):
                    nombre = pairs[j][0].upper() if len(pairs[j]) > 0 else ''
                    funcion = pairs[j][1] if len(pairs[j]) > 1 else ''
                    new_p_el = copy.deepcopy(last_para._p)
                    insert_after_el.addnext(new_p_el)
                    insert_after_el = new_p_el
                    if spacer_el is not None:
                        new_spacer_el = copy.deepcopy(spacer_el)
                        insert_after_el.addnext(new_spacer_el)
                        insert_after_el = new_spacer_el
                    # Wrap the new XML element back into a Paragraph to fill it
                    from docx.text.paragraph import Paragraph
                    new_para = Paragraph(new_p_el, last_para._parent)
                    fill_aditivo_para(new_para, nombre, funcion)

            # Si hay menos aditivos reales que bloques de placeholder fijos
            # en el template, los bloques sobrantes se eliminan en vez de
            # dejarlos con el texto de ejemplo "ADITIVO" / "Función
            # tecnológica" (lo mismo que se ve al clonar de más arriba,
            # pero a la inversa).
            elif len(pairs) < len(aditivo_paras):
                for para in aditivo_paras[len(pairs):]:
                    p_el = para._p
                    siguiente = p_el.getnext()
                    if siguiente is not None and siguiente.tag == qn('w:p'):
                        texto_siguiente = ''.join(
                            t.text or '' for t in siguiente.findall('.//' + qn('w:t'))
                        )
                        if not texto_siguiente.strip() and siguiente.getparent() is not None:
                            siguiente.getparent().remove(siguiente)
                    if p_el.getparent() is not None:
                        p_el.getparent().remove(p_el)

    elif template_id == 'reporte_fibra':
        replace_all(doc, {
            'xxxxxx': producto,
            'xx de febrero del 202x': fecha,
            'xx de febrero del 202X': fecha,
        })
        if data.get('fibra_total'):
            WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            fibra_tbls = doc.element.body.findall(f'.//{{{WNS}}}tbl')
            if fibra_tbls:
                fibra_tbl = fibra_tbls[0]
                rows = fibra_tbl.findall(f'{{{WNS}}}tr')
                if len(rows) > 1:
                    cells = rows[1].findall(f'{{{WNS}}}tc')
                    if len(cells) > 1:
                        set_cell_value_xml(cells[1], str(data['fibra_total']), WNS)

    elif template_id == 'reporte_formula':
        replace_all(doc, {'XXX': producto, 'XX de XX del 2025': fecha})
        if data.get('formula_rows'):
            WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            formula_tbls = doc.element.body.findall(f'.//{{{WNS}}}tbl')
            if formula_tbls:
                tbl_el = formula_tbls[0]
                data_rows = parse_rows(data['formula_rows'])
                # Use the second row (index 1) as the clone template — it's a
                # blank data row without the header's gray shading. Save a
                # deep copy of it BEFORE removing rows, since it gets deleted
                # in the trim loop below.
                all_tr_initial = tbl_el.findall(f'{{{WNS}}}tr')
                template_tr = copy.deepcopy(all_tr_initial[1]) if len(all_tr_initial) > 1 else copy.deepcopy(all_tr_initial[0])
                # Keep first row (header) and last row (TOTAL); remove the rest
                while len(tbl_el.findall(f'{{{WNS}}}tr')) > 2:
                    all_tr = tbl_el.findall(f'{{{WNS}}}tr')
                    tbl_el.remove(all_tr[-2])
                total_tr = tbl_el.findall(f'{{{WNS}}}tr')[-1]
                for row_data in data_rows:
                    new_tr = copy.deepcopy(template_tr)
                    cells = new_tr.findall(f'{{{WNS}}}tc')
                    for i, val in enumerate(row_data):
                        if i < len(cells):
                            val_str = str(val)
                            if i == 1 and val_str.strip():
                                val_str = val_str.strip() + ' %'
                            set_cell_value_xml(cells[i], val_str, WNS)
                    tbl_el.insert(list(tbl_el).index(total_tr), new_tr)

    elif template_id == 'reporte_saborizantes':
        replace_all(doc, {
            'XXXXXXX': producto,
            'XX de febrero del 202X': fecha,
            'XX de febrero del 202x': fecha,
        })
        sab_map = {
            'SABORIZANTE NATURAL': 'sab_natural',
            'SABORIZANTE IDENTICO NATURAL': 'sab_identico',
        }
        WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        sab_tbls = doc.element.body.findall(f'.//{{{WNS}}}tbl')
        if sab_tbls:
            tbl_el = sab_tbls[0]
            filas_a_quitar = []
            for row in tbl_el.findall(f'{{{WNS}}}tr')[1:]:
                cells = row.findall(f'{{{WNS}}}tc')
                if len(cells) < 2: continue
                key = ''.join(t.text for t in cells[0].findall(f'.//{{{WNS}}}t') if t.text).strip()

                fid_de_esta_fila = next((fid for lbl, fid in sab_map.items() if lbl in key), None)
                if fid_de_esta_fila:
                    valor = str(data.get(fid_de_esta_fila, '')).strip()
                    if not valor:
                        # Campo no completado: se quita la fila entera en
                        # vez de dejar un valor de ejemplo por defecto.
                        filas_a_quitar.append(row)
                        continue
                    set_cell_value_xml(cells[1], valor + ' %', WNS)
                elif 'TOTAL' in key:
                    valor_total = str(data.get('sab_total', '')).strip()
                    if not valor_total:
                        filas_a_quitar.append(row)
                        continue
                    val = valor_total + ' %'
                    # Preserve bold on TOTAL row
                    runs = cells[1].findall(f'.//{{{WNS}}}r')
                    if runs:
                        for r in runs:
                            for t_el in r.findall(f'{{{WNS}}}t'): t_el.text = ''
                        t_el = runs[0].find(f'{{{WNS}}}t')
                        if t_el is None:
                            t_el = OxmlElement('w:t')
                            runs[0].append(t_el)
                        t_el.text = val
                    else:
                        set_cell_value_xml(cells[1], val, WNS)
            for row in filas_a_quitar:
                if row.getparent() is not None:
                    row.getparent().remove(row)

    # Apply firmante to ALL templates at the end
    set_producto_bold(doc, producto)
    apply_firmante(doc, data)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── Routes ──────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/api/templates')
def get_templates():
    return jsonify({
        tid: {'label': t['label'], 'fields': t['fields'], 'help': t.get('help', '')}
        for tid, t in TEMPLATES.items()
    })

@app.route('/api/generate', methods=['POST'])
def generate():
    body = request.get_json()
    tid = body.get('template_id')
    data = body.get('data', {})
    if tid not in TEMPLATES:
        return jsonify({'error': 'Template no encontrado'}), 400
    try:
        buf = generate_doc(tid, data)
        label = TEMPLATES[tid]['label'].replace(' ', '_')
        producto = data.get('producto', 'producto').replace(' ', '_')
        fecha_archivo = datetime.today().strftime('%d.%m.%Y')
        filename = f"{fecha_archivo}_{producto}_{label}.docx"
        return send_file(buf, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/api/convert-pdf', methods=['POST'])
def convert_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    f = request.files['file']
    tmp_docx = f'/tmp/upload_{f.filename}'
    tmp_pdf = tmp_docx.replace('.docx', '.pdf')
    f.save(tmp_docx)
    try:
        import shutil
        lo_bin = shutil.which('libreoffice') or shutil.which('soffice') or '/usr/bin/libreoffice'
        subprocess.run([lo_bin, '--headless', '--convert-to', 'pdf', '--outdir', '/tmp', tmp_docx],
                       capture_output=True)
        if os.path.exists(tmp_pdf):
            return send_file(tmp_pdf, as_attachment=True,
                             download_name=f.filename.replace('.docx', '.pdf'),
                             mimetype='application/pdf')
        else:
            return jsonify({'error': 'Conversión fallida. LibreOffice no está disponible en el servidor.'}), 500
    finally:
        if os.path.exists(tmp_docx): os.remove(tmp_docx)

@app.route('/api/debug-lo')
def debug_lo():
    import shutil, subprocess
    result = {}
    result['which_libreoffice'] = shutil.which('libreoffice')
    result['which_soffice'] = shutil.which('soffice')
    try:
        find = subprocess.run(['find', '/', '-name', 'soffice', '-type', 'f'],
                              capture_output=True, text=True, timeout=10)
        result['find_soffice'] = find.stdout.strip().split()
    except Exception as e:
        result['find_error'] = str(e)
    return jsonify(result)

@app.route('/api/import-docx', methods=['POST'])
def import_docx():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    f = request.files['file']
    try:
        doc = Document(f)
        campos = {}
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    campo = row.cells[0].text.strip()
                    valor = row.cells[1].text.strip()
                    if campo and valor and campo != 'campo' and campo != 'Campo':
                        campos[campo] = valor
        if not campos:
            for para in doc.paragraphs:
                txt = para.text.strip()
                for sep in ['|', ':']:
                    if sep in txt:
                        parts = txt.split(sep, 1)
                        if len(parts) == 2:
                            campo = parts[0].strip()
                            valor = parts[1].strip()
                            if campo and valor:
                                campos[campo] = valor
                        break
        return jsonify(campos)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/import-ft', methods=['POST'])
def import_ft():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    f = request.files['file']
    try:
        from ft_extractor import extract_ft
        data = extract_ft(f.read())
        return jsonify(data)
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/api/import-dossier', methods=['POST'])
def import_dossier():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    f = request.files['file']
    try:
        from dossier_extractor import extract_dossier
        data = extract_dossier(f.read())
        return jsonify(data)
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/api/import-metodologias', methods=['POST'])
def import_metodologias():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    f = request.files['file']
    try:
        from metodologia_extractor import extract_metodologias
        mapa = extract_metodologias(f.read())
        return jsonify({'metodologias': mapa})
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500


# ─── Google Drive OAuth routes ────────────────────────────────────────────────

# In-memory temp storage for OAuth flow
_temp_files = {}

@app.route('/api/save-temp', methods=['POST'])
def save_temp():
    import uuid
    f = request.files.get('file')
    producto = request.form.get('producto', 'Sin_producto')
    if not f:
        return jsonify({'error': 'No file'}), 400
    file_id = str(uuid.uuid4())
    _temp_files[file_id] = {
        'bytes': f.read(),
        'filename': f.filename,
        'producto': producto
    }
    return jsonify({'file_id': file_id})

@app.route('/oauth/start')
def oauth_start():
    try:
        from drive import get_auth_url
        file_id = request.args.get('file_id', '')
        url = get_auth_url(state=file_id)
        return jsonify({'url': url})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/oauth/callback')
def oauth_callback():
    try:
        from drive import exchange_code
        from urllib.parse import quote
        code = request.args.get('code')
        file_id = request.args.get('state', '')
        creds = exchange_code(code)
        creds_json = quote(json.dumps(creds))
        return f'<script>window.location.href = "/?creds={creds_json}&file_id={file_id}";</script>'
    except Exception as e:
        return f'<script>window.location.href = "/?drive_error={str(e)}";</script>'

@app.route('/api/upload-drive-temp', methods=['POST'])
def upload_drive_temp():
    try:
        from drive import upload_file
        body = request.get_json()
        file_id = body.get('file_id')
        creds_dict = body.get('creds')
        if not file_id or file_id not in _temp_files:
            return jsonify({'error': 'Archivo temporal no encontrado. Generá el documento de nuevo.'}), 400
        tmp = _temp_files.pop(file_id)
        link = upload_file(creds_dict, tmp['filename'], tmp['bytes'], tmp['producto'])
        return jsonify({'link': link})
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
